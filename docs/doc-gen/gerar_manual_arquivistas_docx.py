from datetime import datetime
from pathlib import Path
import tempfile
import os

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "MANUAL_COMPLETO_ARQUIVISTAS.docx"


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _brand_image_png() -> Path | None:
    """Convert HW1 gradient image to PNG for python-docx compatibility."""
    webp = ROOT / "frontend" / "public" / "branding" / "hw1-gradient.webp"
    if not webp.exists():
        return None

    fd, tmp_name = tempfile.mkstemp(prefix="hw1_gradient_", suffix=".png")
    os.close(fd)
    Path(tmp_name).unlink(missing_ok=True)
    png = Path(tmp_name)
    with Image.open(webp) as img:
        img.convert("RGB").save(png, format="PNG")
    return png


def _save_with_fallback(doc: Document, path: Path) -> Path:
    """Save to default name; if locked by Word, save a timestamped copy."""
    try:
        doc.save(str(path))
        return path
    except PermissionError:
        alt = path.with_name(f"{path.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{path.suffix}")
        doc.save(str(alt))
        return alt


def _placeholder(path: Path, title: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 900), color=(248, 248, 248))
    draw = ImageDraw.Draw(img)
    draw.rectangle((40, 40, 1560, 860), outline=(236, 9, 146), width=6)
    draw.text((90, 90), "HW1 - Captura de Tela", fill=(0, 0, 0))
    draw.text((90, 150), title, fill=(20, 20, 20))
    draw.text((90, 220), "Substitua esta imagem por um print real do sistema.", fill=(80, 80, 80))
    img.save(path, format="PNG")
    return path


def _resolve_screenshot(filename: str, title: str) -> Path:
    candidates = [
        ROOT / "docs" / "screenshots" / filename,
        ROOT / "frontend" / "public" / "screenshots" / filename,
    ]
    for item in candidates:
        if item.exists():
            return item

    return _placeholder(ROOT / "docs" / "screenshots" / filename, title)


def _add_screenshot(doc: Document, caption: str, filename: str) -> None:
    path = _resolve_screenshot(filename, caption)
    doc.add_paragraph(caption)
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(path), width=Inches(6.7))


def main() -> None:
    doc = Document()

    brand = _brand_image_png()
    if brand:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(brand), width=Inches(6.2))

    add_title(doc, "Manual Completo do Sistema - Gestao de Apoio Arquivistico")
    add_p(doc, "Publico-alvo: Arquivistas, Gestores de Documentos, Classificadores e Auditores.")
    add_p(doc, f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    add_h1(doc, "1. Visao Geral")
    add_p(doc, "O sistema organiza o ciclo documental de ponta a ponta: entrevistas, classificacao, temporalidade, governanca, integracao, migracao e rastreabilidade.")
    add_bullets(doc, [
        "Entrevistas assistidas para coleta e classificacao inicial.",
        "PCD (Plano de Classificacao Documental) por niveis.",
        "TTD (Tabela de Temporalidade) com regras, legal hold e ordens.",
        "Governanca e trilha de auditoria com verificacao de integridade.",
        "Integracao de acervo e dados de migracao (inventarios, cleansing e ondas).",
        "Ciclo de vida com jobs de retencao e selos de evidencia.",
        "Portal do cliente para resposta externa de entrevistas.",
        "Relatorios e exportacoes PDF/CSV/Excel para operacao.",
    ])

    add_h1(doc, "2. Perfis e Acessos")
    add_h2(doc, "2.1 Perfis Internos")
    add_bullets(doc, [
        "admin: administracao geral, usuarios e configuracoes.",
        "gestor: supervisao operacional dos modulos.",
        "arquivista/classificador: operacao principal de entrevistas, PCD e TTD.",
        "auditor/viewer: consulta e auditoria com menor permissao.",
    ])
    add_h2(doc, "2.2 Perfil Cliente")
    add_bullets(doc, [
        "Acessa somente /portal e entrevistas atribuidas.",
        "Pode editar respostas quando status = em_andamento ou devolvida.",
        "Submete entrevista para analise interna (status submetida).",
    ])

    add_h1(doc, "3. Primeiros Passos")
    add_h2(doc, "3.1 Acesso")
    add_bullets(doc, [
        "Tela de login em /login.",
        "Usuarios internos sao direcionados para /dashboard.",
        "Usuarios cliente sao direcionados para /portal.",
    ])
    add_h2(doc, "3.2 Troca de Senha")
    add_bullets(doc, [
        "No dashboard, acesse Perfil.",
        "Informe senha atual, nova senha e confirmacao.",
        "Regra minima: 8 caracteres, 1 letra maiuscula e 1 numero.",
    ])

    add_h1(doc, "4. Modulo Entrevistas")
    add_h2(doc, "4.1 Fluxo Interno")
    add_bullets(doc, [
        "Criar roteiro (titulo, area, descricao).",
        "Adicionar perguntas e condicoes de exibicao.",
        "Iniciar entrevista e opcionalmente atribuir cliente.",
        "Gerenciar status e evidencias por entrevista.",
    ])
    add_h2(doc, "4.2 Estados da Entrevista")
    add_bullets(doc, [
        "em_andamento: preenchimento ativo.",
        "submetida: cliente finalizou e enviou para revisao.",
        "devolvida: interno devolveu para ajuste com motivo.",
        "concluida: validada e finalizada.",
        "cancelada: encerrada sem conclusao.",
    ])

    add_h1(doc, "5. PCD - Plano de Classificacao")
    add_bullets(doc, [
        "Cadastra niveis hierarquicos: funcao, subfuncao, atividade, serie, classe e tipo_documental.",
        "Cada nivel tem codigo unico, titulo e metadados de sigilo.",
        "Permite versao e manutencao progressiva da arvore documental.",
    ])

    add_h1(doc, "6. TTD - Temporalidade")
    add_bullets(doc, [
        "Criar e atualizar regras de retencao por nivel do PCD.",
        "Destinacoes: eliminacao, guarda_permanente, microfilmagem, amostragem.",
        "Aplicar e revogar legal holds.",
        "Emitir ordens de destinacao quando aplicavel.",
    ])

    add_h1(doc, "7. Governanca e Auditoria")
    add_bullets(doc, [
        "Matriz de rastreabilidade entre norma e classificacao documental.",
        "CRUD da matriz com log de auditoria encadeado por hash.",
        "Consulta de logs e verificacao de integridade.",
    ])

    add_h1(doc, "8. Integracao e Migracao")
    add_h2(doc, "8.1 Integracao")
    add_bullets(doc, [
        "Importacao de acervo (CSV).",
        "Reprocessamento de importacoes.",
        "Exclusao logica de importacoes sem sucesso persistido.",
    ])
    add_h2(doc, "8.2 Dados e Migracao")
    add_bullets(doc, [
        "Regras de cleansing com CRUD.",
        "Inventario de qualidade e indicadores.",
        "Planejamento por ondas com validacao e rollback.",
    ])

    add_h1(doc, "9. Ciclo de Vida")
    add_bullets(doc, [
        "Criacao e execucao de jobs de retencao.",
        "Cancelamento/exclusao de jobs conforme status.",
        "Selos de evidencia e pacote de auditoria.",
        "Exclusao de selo apenas em rascunho.",
    ])

    add_h1(doc, "10. Conhecimento")
    add_bullets(doc, [
        "Templates e guias com persistencia em banco.",
        "CRUD de templates e trilhas de onboarding.",
        "Controle de progresso por usuario.",
    ])

    add_h1(doc, "11. Relatorios e Exportacoes")
    add_bullets(doc, [
        "Busca avancada no PCD.",
        "Dashboard de temporalidade.",
        "Exportacao de CCD/TTD em PDF, CSV e Excel.",
        "Termo de eliminacao e relatorio de transferencia/recolhimento (inclusive PDF).",
        "Importacao em lote via Excel para aceleracao de cadastro.",
    ])

    add_h1(doc, "12. Portal do Cliente")
    add_bullets(doc, [
        "Lista entrevistas atribuidas ao cliente autenticado.",
        "Edicao de respostas e evidencias no escopo permitido.",
        "Submissao e acompanhamento de status.",
        "Bloqueio automatico de acesso do cliente aos modulos internos.",
    ])

    add_h1(doc, "13. Operacao Local (Docker)")
    add_bullets(doc, [
        "Subida: docker compose up -d --build",
        "Status: docker compose ps",
        "Logs: docker compose logs backend/frontend/celery-worker --tail 200",
        "Health backend: GET /health na porta 8000",
        "Frontend: porta 4000 (primeira compilacao pode demorar)",
    ])

    add_h1(doc, "14. Troubleshooting")
    add_h2(doc, "14.1 Frontend lento ao iniciar")
    add_p(doc, "Causa comum: primeira compilacao do Next.js em modo desenvolvimento. Aguarde a mensagem 'Ready' no log do frontend antes do primeiro acesso.")
    add_h2(doc, "14.2 Celery warning de root")
    add_p(doc, "Causa: worker executado com usuario root no container de desenvolvimento. Nao bloqueia operacao local, mas recomenda-se ajustar user no Dockerfile/compose para hardening de producao.")
    add_h2(doc, "14.3 Endpoint /api/v1/health")
    add_p(doc, "No backend atual o health padrao exposto esta em /health.")

    add_h1(doc, "15. Boas Praticas para Arquivistas")
    add_bullets(doc, [
        "Padronizar codigos e nomenclaturas no PCD antes de criar regras no TTD.",
        "Registrar base legal de cada regra de retencao.",
        "Usar devolucao com motivo claro e objetivo para orientar clientes.",
        "Anexar evidencias e revisar logs de auditoria em processos criticos.",
        "Validar relatorios periodicamente para detectar gargalos.",
    ])

    add_h1(doc, "16. Anexo - Checklist Operacional Diario")
    add_bullets(doc, [
        "Verificar /health no inicio do expediente.",
        "Checar jobs de retencao agendados e pendencias.",
        "Acompanhar entrevistas devolvidas/submetidas.",
        "Monitorar integracoes com erro e acionar reprocessamento.",
        "Emitir relatorios de apoio para controle gerencial.",
    ])

    add_h1(doc, "17. Capturas de Tela")
    add_p(doc, "Abaixo estao os prints do sistema para apoio de treinamento. Se algum print ainda nao existir, o documento inclui um placeholder visual para substituicao.")
    _add_screenshot(doc, "Tela de Login", "login.png")
    _add_screenshot(doc, "Portal do Cliente", "portal.png")
    _add_screenshot(doc, "Dashboard Administrativo", "dashboard.png")
    _add_screenshot(doc, "Swagger / Documentacao da API", "api_docs.png")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    saved = _save_with_fallback(doc, OUT)
    if brand and brand.exists():
        brand.unlink(missing_ok=True)
    print(f"Manual gerado em: {saved}")


if __name__ == "__main__":
    main()
