from datetime import datetime
from pathlib import Path
import tempfile
import os

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "MANUAL_EXECUTIVO_TREINAMENTO_ARQUIVISTAS.docx"


def title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _brand_image_png() -> Path | None:
    webp = ROOT / "frontend" / "public" / "branding" / "hw1-gradient.webp"
    if not webp.exists():
        return None

    fd, tmp_name = tempfile.mkstemp(prefix="hw1_gradient_exec_", suffix=".png")
    os.close(fd)
    Path(tmp_name).unlink(missing_ok=True)
    png = Path(tmp_name)
    with Image.open(webp) as img:
        img.convert("RGB").save(png, format="PNG")
    return png


def _save_with_fallback(doc: Document, path: Path) -> Path:
    try:
        doc.save(str(path))
        return path
    except PermissionError:
        alt = path.with_name(f"{path.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{path.suffix}")
        doc.save(str(alt))
        return alt


def _placeholder(path: Path, title: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1400, 800), color=(250, 250, 250))
    draw = ImageDraw.Draw(img)
    draw.rectangle((30, 30, 1370, 770), outline=(0, 178, 168), width=5)
    draw.text((80, 80), "HW1 - Print de Treinamento", fill=(0, 0, 0))
    draw.text((80, 140), title, fill=(20, 20, 20))
    draw.text((80, 200), "Substitua por captura real para a turma.", fill=(90, 90, 90))
    img.save(path, format="PNG")
    return path


def _resolve_screenshot(filename: str, title: str) -> Path:
    candidates = [
        ROOT / "docs" / "screenshots" / filename,
        ROOT / "frontend" / "public" / "screenshots" / filename,
    ]
    for item in candidates:
        if item.exists():
            return item
    return _placeholder(ROOT / "docs" / "screenshots" / filename, title)


def _add_screenshot(doc: Document, caption: str, filename: str) -> None:
    path = _resolve_screenshot(filename, caption)
    doc.add_paragraph(caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.5))


def main() -> None:
    doc = Document()

    brand = _brand_image_png()
    if brand:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(brand), width=Inches(6.0))

    title(doc, "Manual Executivo de Treinamento - Arquivistas")
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("Objetivo: capacitar rapidamente a equipe para operar o sistema no dia a dia.")

    doc.add_heading("1. Fluxo Diario (15 min)", level=1)
    bullets(doc, [
        "Acessar o dashboard e conferir status geral dos modulos.",
        "Validar health do backend e fila de jobs.",
        "Verificar entrevistas submetidas e devolvidas.",
        "Monitorar importacoes com erro e acionar reprocessamento.",
    ])

    doc.add_heading("2. Roteiro Operacional (30 min)", level=1)
    bullets(doc, [
        "Criar/ajustar roteiro de entrevista e perguntas.",
        "Atribuir entrevista para cliente quando necessario.",
        "Revisar respostas e evidencias anexadas.",
        "Concluir ou devolver com motivo objetivo.",
    ])

    doc.add_heading("3. Classificacao e Temporalidade (30 min)", level=1)
    bullets(doc, [
        "Manter arvore PCD atualizada (funcao ate tipo documental).",
        "Criar e revisar regras TTD por nivel.",
        "Aplicar hold quando houver impedimento legal.",
        "Emitir ordens de destinacao conforme politica.",
    ])

    doc.add_heading("4. Governanca e Auditoria (20 min)", level=1)
    bullets(doc, [
        "Atualizar matriz de rastreabilidade (norma x classe).",
        "Consultar logs de auditoria e integridade.",
        "Registrar evidencias de decisao em processos criticos.",
    ])

    doc.add_heading("5. Relatorios Essenciais (20 min)", level=1)
    bullets(doc, [
        "Gerar dashboard de temporalidade para visao executiva.",
        "Exportar PDF/CSV/Excel para reunioes e compliance.",
        "Emitir termo de eliminacao e relatorio de transferencia/recolhimento.",
    ])

    doc.add_heading("6. Atalhos de Suporte", level=1)
    bullets(doc, [
        "Se frontend estiver lento no primeiro acesso, aguarde compilacao inicial do Next.js.",
        "Health backend principal: /health.",
        "Logs locais: docker compose logs backend/frontend/celery-worker --tail 200.",
    ])

    doc.add_heading("7. Checklist de Encerramento (5 min)", level=1)
    bullets(doc, [
        "Pendencias de entrevistas tratadas.",
        "Jobs de retencao sem erro.",
        "Importacoes com falha registradas para acao.",
        "Relatorio diario exportado quando aplicavel.",
    ])

    doc.add_heading("8. Prints de Referencia", level=1)
    _add_screenshot(doc, "Login", "login.png")
    _add_screenshot(doc, "Dashboard", "dashboard.png")
    _add_screenshot(doc, "API Docs", "api_docs.png")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    saved = _save_with_fallback(doc, OUT)
    if brand and brand.exists():
        brand.unlink(missing_ok=True)
    print(f"Manual executivo gerado em: {saved}")


if __name__ == "__main__":
    main()
